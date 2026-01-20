"""
Markdown to Word document (.docx) converter for AI-Ready File Converter.
Converts Markdown files to Word documents with full formatting support.
"""

import re
from io import BytesIO
from pathlib import Path
from typing import Any, List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import ReverseConverter


class MarkdownToDocxConverter(ReverseConverter):
    """Converter for Markdown files to Word documents."""
    
    def _get_source_type(self) -> str:
        return "markdown"
    
    def _get_target_type(self) -> str:
        return "word_document"
    
    def _get_target_extension(self) -> str:
        return ".docx"
    
    def _parse_source(self) -> str:
        """Read and return the markdown content."""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _generate_output(self, markdown_content: str) -> bytes:
        """Convert markdown content to a Word document."""
        doc = Document()
        
        # Set up styles for code blocks
        self._setup_styles(doc)
        
        # Parse and convert markdown
        self._convert_markdown_to_doc(doc, markdown_content)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _setup_styles(self, doc: Document):
        """Set up custom styles for the document."""
        styles = doc.styles
        
        # Create a code block style if it doesn't exist
        try:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.left_indent = Inches(0.25)
            # Add shading
            self._add_shading(code_style.element, "E8E8E8")
        except ValueError:
            # Style already exists
            pass
        
        # Create inline code style
        try:
            inline_code = styles.add_style('InlineCode', WD_STYLE_TYPE.CHARACTER)
            inline_code.font.name = 'Courier New'
            inline_code.font.size = Pt(10)
        except ValueError:
            pass
    
    def _add_shading(self, element, color: str):
        """Add background shading to an element."""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        element.get_or_add_pPr().append(shd)
    
    def _convert_markdown_to_doc(self, doc: Document, content: str):
        """Parse markdown and add elements to the document."""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Skip the header block if it looks like our converter's output
            if line.startswith('# ') and i == 0:
                # Check if this is a converted file header
                if i + 1 < len(lines) and lines[i + 1].startswith('> Converted from'):
                    # Skip the header block (title, converted note, separator)
                    i += 1
                    while i < len(lines) and lines[i] != '---':
                        i += 1
                    i += 1  # Skip the ---
                    if i < len(lines) and lines[i].strip() == '':
                        i += 1  # Skip blank line after ---
                    continue
            
            # Code block (fenced)
            if line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # Skip closing ```
                
                # Add code block
                for code_line in code_lines:
                    p = doc.add_paragraph(code_line, style='CodeBlock')
                continue
            
            # Table
            if line.startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                self._add_table(doc, table_lines)
                continue
            
            # Heading
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Word supports Heading 1-9
                heading_style = f'Heading {min(level, 9)}'
                p = doc.add_paragraph(style=heading_style)
                self._add_formatted_text(p, text)
                i += 1
                continue
            
            # Horizontal rule
            if re.match(r'^(-{3,}|\*{3,}|_{3,})$', line.strip()):
                # Add a horizontal line using a paragraph border
                p = doc.add_paragraph()
                self._add_horizontal_rule(p)
                i += 1
                continue
            
            # Unordered list
            if re.match(r'^[\s]*[-*+]\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                    match = re.match(r'^([\s]*)[-*+]\s+(.+)$', lines[i])
                    if match:
                        indent = len(match.group(1)) // 2
                        list_items.append((indent, match.group(2)))
                    i += 1
                
                for indent, item_text in list_items:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
                    self._add_formatted_text(p, item_text)
                continue
            
            # Ordered list
            if re.match(r'^[\s]*\d+\.\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                    match = re.match(r'^([\s]*)\d+\.\s+(.+)$', lines[i])
                    if match:
                        indent = len(match.group(1)) // 2
                        list_items.append((indent, match.group(2)))
                    i += 1
                
                for indent, item_text in list_items:
                    p = doc.add_paragraph(style='List Number')
                    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
                    self._add_formatted_text(p, item_text)
                continue
            
            # Blockquote
            if line.startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].startswith('>'):
                    quote_text = re.sub(r'^>\s*', '', lines[i])
                    quote_lines.append(quote_text)
                    i += 1
                
                for quote_line in quote_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                    # Add italic formatting for quotes
                    run = p.add_run(quote_line)
                    run.italic = True
                continue
            
            # Empty line
            if line.strip() == '':
                i += 1
                continue
            
            # Regular paragraph
            p = doc.add_paragraph()
            self._add_formatted_text(p, line)
            i += 1
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code, links)."""
        # Pattern to match formatting: **bold**, *italic*, `code`, [text](url)
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            full_match = match.group(0)
            
            # Bold
            if full_match.startswith('**') and full_match.endswith('**'):
                run = paragraph.add_run(match.group(2))
                run.bold = True
            # Italic
            elif full_match.startswith('*') and full_match.endswith('*'):
                run = paragraph.add_run(match.group(3))
                run.italic = True
            # Inline code
            elif full_match.startswith('`') and full_match.endswith('`'):
                run = paragraph.add_run(match.group(4))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Link
            elif full_match.startswith('['):
                link_text = match.group(5)
                link_url = match.group(6)
                self._add_hyperlink(paragraph, link_url, link_text)
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add a hyperlink to a paragraph."""
        # Get the document part
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a run element for the text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add blue color and underline
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        
        # Add the text
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _add_table(self, doc: Document, table_lines: List[str]):
        """Add a markdown table to the document."""
        if len(table_lines) < 2:
            return
        
        # Parse header row
        header_cells = self._parse_table_row(table_lines[0])
        if not header_cells:
            return
        
        # Skip separator row (line with dashes)
        data_start = 1
        if len(table_lines) > 1 and re.match(r'^\|[\s\-:|]+\|$', table_lines[1]):
            data_start = 2
        
        # Parse data rows
        data_rows = []
        for line in table_lines[data_start:]:
            cells = self._parse_table_row(line)
            if cells:
                data_rows.append(cells)
        
        # Create table
        num_cols = len(header_cells)
        num_rows = 1 + len(data_rows)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Add header
        header_row = table.rows[0]
        for i, cell_text in enumerate(header_cells):
            if i < len(header_row.cells):
                cell = header_row.cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(data_rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.paragraphs[0].add_run(cell_text)
        
        # Add blank paragraph after table
        doc.add_paragraph()
    
    def _parse_table_row(self, line: str) -> List[str]:
        """Parse a markdown table row into cells."""
        # Remove leading/trailing pipes and split
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        
        cells = [cell.strip() for cell in line.split('|')]
        return cells
    
    def _add_horizontal_rule(self, paragraph):
        """Add a horizontal rule (bottom border) to a paragraph."""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
