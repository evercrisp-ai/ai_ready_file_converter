"""
Word document (.docx) converter for Human-to-AI Format Converter.
Extracts text, headings, paragraphs, and tables from Word documents.
"""

from pathlib import Path
from docx import Document
from docx.table import Table
from .base import BaseConverter

# #region agent log
import json as _json
from datetime import datetime as _dt
_LOG_PATH = "/Users/benscooper/human_format_to_ai_converter/.cursor/debug.log"
def _dbg(loc, msg, data, hyp):
    with open(_LOG_PATH, "a") as f:
        f.write(_json.dumps({"location": loc, "message": msg, "data": data, "hypothesisId": hyp, "timestamp": _dt.now().isoformat(), "sessionId": "debug-session"}) + "\n")
# #endregion


class DocxConverter(BaseConverter):
    """Converter for Microsoft Word documents."""
    
    def _get_file_type(self) -> str:
        return "word_document"
    
    def _extract_content(self) -> dict:
        """Extract content from Word document."""
        # #region agent log
        _dbg("docx_converter.py:28", "Starting _extract_content", {"file_path": str(self.file_path)}, "A")
        # #endregion
        try:
            doc = Document(self.file_path)
            # #region agent log
            _dbg("docx_converter.py:33", "Document loaded successfully", {"para_count": len(doc.paragraphs), "table_count": len(doc.tables)}, "A")
            # #endregion
        except Exception as e:
            # #region agent log
            _dbg("docx_converter.py:37", "Document loading FAILED", {"error": str(e), "error_type": type(e).__name__}, "A")
            # #endregion
            raise
        
        content = {
            "text": "",
            "paragraphs": [],
            "tables": [],
            "headings": []
        }
        
        full_text_parts = []
        
        # #region agent log
        body_elements = list(doc.element.body) if doc.element.body is not None else []
        _dbg("docx_converter.py:48", "Body elements count", {"count": len(body_elements), "body_is_none": doc.element.body is None}, "B")
        # #endregion
        
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                # #region agent log
                matched = False
                # #endregion
                for para in doc.paragraphs:
                    if para._element == element:
                        # #region agent log
                        matched = True
                        # #endregion
                        text = para.text.strip()
                        if text:
                            # #region agent log
                            try:
                                style_name = para.style.name if para.style else ""
                            except AttributeError as e:
                                _dbg("docx_converter.py:68", "Style attribute error", {"error": str(e), "para_text": text[:50]}, "D")
                                style_name = ""
                            # #endregion
                            
                            # Check if it's a heading
                            if style_name.startswith('Heading'):
                                try:
                                    level = int(style_name.replace('Heading ', ''))
                                except ValueError:
                                    level = 1
                                content["headings"].append({
                                    "level": level,
                                    "text": text
                                })
                            
                            content["paragraphs"].append({
                                "text": text,
                                "style": style_name
                            })
                            full_text_parts.append(text)
                        break
                # #region agent log
                if not matched:
                    _dbg("docx_converter.py:91", "Paragraph element NOT matched", {"tag": str(element.tag)}, "B")
                # #endregion
            
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                # #region agent log
                table_matched = False
                # #endregion
                for table in doc.tables:
                    if table._element == element:
                        # #region agent log
                        table_matched = True
                        # #endregion
                        try:
                            table_data = self._extract_table(table)
                            content["tables"].append(table_data)
                        except Exception as e:
                            # #region agent log
                            _dbg("docx_converter.py:109", "Table extraction FAILED", {"error": str(e)}, "E")
                            # #endregion
                        break
                # #region agent log
                if not table_matched:
                    _dbg("docx_converter.py:114", "Table element NOT matched", {"tag": str(element.tag)}, "E")
                # #endregion
        
        content["text"] = "\n\n".join(full_text_parts)
        
        # #region agent log
        _dbg("docx_converter.py:120", "Content extraction complete", {
            "paragraphs_extracted": len(content["paragraphs"]),
            "tables_extracted": len(content["tables"]),
            "headings_extracted": len(content["headings"]),
            "text_length": len(content["text"])
        }, "C")
        # #endregion
        
        # Update metadata
        self._metadata = {
            "paragraph_count": len(content["paragraphs"]),
            "table_count": len(content["tables"]),
            "heading_count": len(content["headings"]),
            "word_count": len(content["text"].split())
        }
        
        return content
    
    def _extract_table(self, table: Table) -> dict:
        """Extract data from a Word table."""
        rows = []
        headers = []
        
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = row_data
            rows.append(row_data)
        
        return {
            "headers": headers,
            "rows": rows
        }
    
    def _format_as_markdown(self, content: dict) -> str:
        """Format Word document content as Markdown."""
        md_parts = []
        
        # Process paragraphs with proper heading formatting
        for para in content["paragraphs"]:
            text = para["text"]
            style = para.get("style", "")
            
            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading ", ""))
                    level = min(level, 6)  # Markdown supports h1-h6
                except ValueError:
                    level = 1
                md_parts.append(f"{'#' * level} {text}")
            else:
                md_parts.append(text)
        
        # Add tables
        for i, table in enumerate(content["tables"]):
            if table["rows"]:
                md_parts.append("")  # Empty line before table
                
                # Header row
                if table["headers"]:
                    md_parts.append("| " + " | ".join(table["headers"]) + " |")
                    md_parts.append("| " + " | ".join(["---"] * len(table["headers"])) + " |")
                
                # Data rows (skip first if it was headers)
                start_idx = 1 if table["headers"] else 0
                for row in table["rows"][start_idx:]:
                    md_parts.append("| " + " | ".join(row) + " |")
                
                md_parts.append("")  # Empty line after table
        
        return "\n\n".join(md_parts)
